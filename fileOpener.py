from textual import on, work
from textual.app import App, ComposeResult
from textual.containers import HorizontalGroup, VerticalScroll, Container
from textual.widgets import Button, Label, Header, Footer, LoadingIndicator, Input, Log

from textual_fspicker import FileOpen

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph





class ResumeAtsKiller(App[None]):
    
    
    #CSS_PATH = "app.tcss"
    JobDescriptionURL = None
    GotNewSkillsList = False
    
    
    CSS_PATH = "app.tcss"
    
    def compose(self) -> ComposeResult:
        yield Input(placeholder='Job Description URL Here...', type="text", id="JobURL")
        yield Button("Submit URL", id="FetchJobDescription", variant="success")
        yield Button("Select your resume", id="openFile", variant="primary",)
        yield Label()
        yield Header()
        
        
    def add_custom_bullet_style(self, doc):
        styles = doc.styles
        
        # Check if the style already exists to avoid duplication
        if 'CustomBullet' not in [s.name for s in styles]:
            bullet_style = styles.add_style('CustomBullet', WD_STYLE_TYPE.PARAGRAPH)
            bullet_style.font.name = 'Arial'
            bullet_style.font.size = Pt(12)
            bullet_style.paragraph_format.left_indent = Pt(18)
            bullet_style.paragraph_format.space_before = Pt(6)
            bullet_style.paragraph_format.space_after = Pt(6)

            # Add bullet character (Unicode bullet) manually
            bullet_style.paragraph_format.first_line_indent = Pt(-18)
            
 
    # so we can insert custom text after a specific paragraph...
    def insert_paragraph_after(self, paragraph, text, style=None, formatDict=None):
        
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        
        
        # code to try and copy style...
        #if style is not None:
        #    new_para.style = style
        
        # try and apply formatting...
        return new_para

  
        
    def get_paragraph_format(self, source_para):
        format_dict = {}
        
        # Capture paragraph-level formatting
        format_dict['style'] = source_para.style.name
        
        # Capture run-level formatting (from the first run, if available)
        if source_para.runs:
            run = source_para.runs[0]
            format_dict['bold'] = run.bold
            format_dict['italic'] = run.italic
            format_dict['underline'] = run.underline
            format_dict['font_size'] = run.font.size
            format_dict['font_name'] = run.font.name
            format_dict['font_color'] = run.font.color.rgb if run.font.color else None
        
        # Capture list formatting (if part of a list)
        pPr = source_para._element.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            format_dict['list_level'] = numPr.find(qn('w:ilvl')).get(qn('w:val')) if numPr.find(qn('w:ilvl')) else None
            format_dict['list_id'] = numPr.find(qn('w:numId')).get(qn('w:val')) if numPr.find(qn('w:numId')) else None

        return format_dict
    
    
    
    def apply_paragraph_format(self, target_para, format_dict):
        # Apply style (if it exists in the document's style set)
        if format_dict.get('style'):
            try:
                target_para.style = format_dict['style']
            except KeyError:
                print(f"Style '{format_dict['style']}' not found. Skipping style.")

        # Add a run and apply formatting directly to it
        if target_para.runs:
            run = target_para.runs[0]
        else:
            run = target_para.add_run()
        
        run.bold = format_dict.get('bold')
        run.italic = format_dict.get('italic')
        run.underline = format_dict.get('underline')
        
        # Handle font settings carefully (check for None)
        if format_dict.get('font_size') is not None:
            run.font.size = format_dict['font_size']
        if format_dict.get('font_name'):
            run.font.name = format_dict['font_name']
        if format_dict.get('font_color'):
            run.font.color.rgb = format_dict['font_color']

        # Apply list formatting (if available)
        if format_dict.get('list_id') and format_dict.get('list_level') is not None:
            pPr = target_para._element.get_or_add_pPr()
            numPr = OxmlElement('w:numPr')
            
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), format_dict['list_level'])
            numPr.append(ilvl)
            
            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), format_dict['list_id'])
            numPr.append(numId)
            
            pPr.append(numPr)

        # Ensure consistent alignment and spacing
        para_format = target_para.paragraph_format
        para_format.left_indent = format_dict.get('left_indent')
        para_format.right_indent = format_dict.get('right_indent')
        para_format.space_before = format_dict.get('space_before')
        para_format.space_after = format_dict.get('space_after')
        para_format.line_spacing = format_dict.get('line_spacing')
     
        """   # Apply paragraph-level formatting
            target_para.style = format_dict.get('style')
            
            
            
            # Apply run-level formatting
            run = target_para.add_run()
            run.bold = format_dict.get('bold')
            run.italic = format_dict.get('italic')
            run.underline = format_dict.get('underline')
            run.font.size = format_dict.get('font_size')
            run.font.name = format_dict.get('font_name')
            run.font.color.rgb = format_dict.get('font_color')
            
            # Apply list formatting (if available)
            if format_dict.get('list_id') and format_dict.get('list_level') is not None:
                pPr = target_para._element.get_or_add_pPr()
                numPr = OxmlElement('w:numPr')
                
                ilvl = OxmlElement('w:ilvl')
                ilvl.set(qn('w:val'), format_dict['list_level'])
                numPr.append(ilvl)
                
                numId = OxmlElement('w:numId')
                numId.set(qn('w:val'), format_dict['list_id'])
                numPr.append(numId)
                
                pPr.append(numPr)
        """
    
    
    def list_number(self, paragraph, level):
        pPr = paragraph._element.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(level))
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')  # Use list ID 1
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)
    
    def makeLog(self, textToWrite) -> ComposeResult:
        
        log = self.query_one(Log)
        log.write_line(textToWrite)
        
    def parseResume(self, resumePath) -> None:
        
        # make sure we only try and parse .docx files. save us and the user a little headache.
        if not (str(resumePath).lower().endswith('.docx')):
            self.notify("Error: Please select a word document, or use the included converter for PDFs.", severity="error")
        else:
            self.notify("Attempting to open: " + str(resumePath))
            
            
            OpenError = False
            try:
                usersOldResume = Document(str(resumePath))
            except Exception as ResumeOpenError:
                # if, by some miracle, the user tries to open a fake .docx, or one that is too old, we'll catch that here.
                self.notify("Error while trying to open your resume: " + str(ResumeOpenError), severity="error", timeout=10)
                OpenError = True
            
            if not OpenError:
                # here is where we would begin going thru the user's resume.
                # try looking for section headers or something that start with "Skills"
                
                
                FoundSkillSection = False
                
                

                # add the bullet style in case the user has that one weird issue where you cant check the doc style
                self.add_custom_bullet_style(usersOldResume)
                
                skillsParagraph = None
                previousFormatting = None
                
               
                for para in usersOldResume.paragraphs:
                    # grab formatting from template...
                    if '[skills]' in para.text.lower() and previousFormatting == None:
                        self.notify("captured formatting...")
                        previousFormatting = self.get_paragraph_format(para)
                        self.notify(str(previousFormatting), timeout=20)
                    if 'skills' in para.text.strip().lower() and para.style.name.startswith('Heading'):
                        FoundSkillSection = True
                        skillsParagraph = para
                        is_in_skills_section = True
                        continue
                     # remove the old skills paragraph....
                    if is_in_skills_section:
                        if para.text.strip() == '' or para.style.name.startswith('Heading'):
                            is_in_skills_section = False
                        else:
                            para.clear()
                
                
                # add new skills.
                new_skills = ['Python', "Malware Analysis", "Hotdog Eating contests", "potatoes"]
                
                
                if FoundSkillSection:
                    for skill in new_skills:
                        bullet = f' {skill}'
                        #self.insert_paragraph_after(skillsParagraph, bullet, style='List Bullet' if 'List Bullet' in [s.name for s in usersOldResume.styles] else None)
                        new_skills_para = self.insert_paragraph_after(skillsParagraph, bullet, style='List Paragraph' if 'List Paragraph' in [s.name for s in usersOldResume.styles] else None, formatDict=previousFormatting)
                        self.apply_paragraph_format(new_skills_para, previousFormatting)
                    self.notify("New Skills Added!")
                
                
                
                
                # save changes to a new document, we do NOT want to actually destroy anything.
                usersOldResume.save("./test5.docx")
                self.notify("Changes Saved to a new copy of resume at: ")
                
    #@on(Button.Pressed)
    #@work
    
    @on(Input.Changed)
    def saveURL(self, event: Input.Changed) -> None:
        JobDescriptionURL = Input._value
    @work
    async def open_a_file(self) -> None:
        if opened := await self.push_screen_wait(FileOpen()):
            self.parseResume(opened)
    
    def on_button_pressed(self, event: Button.Pressed) -> None:
        # only let the user upload their resume once they've entered a job site description...
        if event.button.id == "FetchJobDescription":
            #self.notify(JobDescriptionURL)
            self.JobDescriptionURL= self.query_one("#JobURL", Input).value
            self.notify(self.JobDescriptionURL)
        if event.button.id == "openFile" and self.JobDescriptionURL is not None:
            self.open_a_file()
        if event.button.id == "openFile" and self.JobDescriptionURL is None: self.notify('Must submit job description URL first!', severity='error')
            


if __name__ == "__main__":
    ResumeAtsKiller().run()